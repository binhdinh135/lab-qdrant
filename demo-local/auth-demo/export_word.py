"""
Script convert báo cáo markdown sang file Word (.docx).

Cách dùng:
  D:\\Qdrant\\.venv\\Scripts\\pip.exe install python-docx
  D:\\Qdrant\\.venv\\Scripts\\python.exe export_word.py

Output: BAO_CAO_PHAN_QUYEN_QDRANT.docx
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Cần cài python-docx:")
    print("  D:\\Qdrant\\.venv\\Scripts\\pip.exe install python-docx")
    exit(1)


def add_code_block(doc, text):
    """Thêm code block vào document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_table(lines):
    """Parse markdown table thành list of rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and not line.startswith("|---") and not line.startswith("|-"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells and not all(c.startswith("-") for c in cells):
                rows.append(cells)
    return rows


def add_table(doc, rows):
    """Thêm table vào document."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            if i == 0:  # Header row bold
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def main():
    script_dir = Path(__file__).resolve().parent
    md_path = script_dir / "BAO_CAO_PHAN_QUYEN_QDRANT.md"
    output_path = script_dir / "BAO_CAO_PHAN_QUYEN_QDRANT.docx"

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    for line in lines:
        # Code block
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Flush table if pending
                if in_table and table_lines:
                    rows = parse_table(table_lines)
                    add_table(doc, rows)
                    table_lines = []
                    in_table = False
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Table
        if line.strip().startswith("|"):
            in_table = True
            table_lines.append(line)
            continue
        elif in_table:
            # End of table
            rows = parse_table(table_lines)
            add_table(doc, rows)
            table_lines = []
            in_table = False

        # Headers
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("---"):
            continue
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.bold = True
        elif line.strip().startswith(">"):
            p = doc.add_paragraph(line.strip().lstrip("> "), style="Quote")
        elif line.strip():
            doc.add_paragraph(line.strip())

    # Flush remaining table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        add_table(doc, rows)

    doc.save(str(output_path))
    print(f"✅ Đã tạo file Word: {output_path}")
    print(f"   Kích thước: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
